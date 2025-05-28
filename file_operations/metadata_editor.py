#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文件元数据编辑器

这个脚本提供了查看和修改各种文件类型元数据的功能，包括图像EXIF数据、
音频ID3标签、视频元数据和文档属性等。支持批量处理文件和多种格式的
元数据导入导出，适用于摄影作品整理、媒体文件分类和数据清理等场景。
"""

import argparse
import csv
import fnmatch
import json
import logging
import os
import re
import sys
from enum import Enum
from typing import Dict, List, Any, Optional

# 设置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

# 尝试导入处理各种文件类型所需的库
DEPENDENCIES = {
    "image": False,
    "audio": False,
    "video": False,
    "document": False
}

try:
    from PIL import Image, ExifTags
    import piexif

    DEPENDENCIES["image"] = True
except ImportError:
    logger.warning("图像处理库缺失。安装PIL/Pillow和piexif以支持图像元数据: pip install Pillow piexif")

try:
    import mutagen
    from mutagen.id3 import ID3, APIC, TIT2, TPE1, TALB, TDRC
    from mutagen.flac import FLAC
    from mutagen.mp3 import MP3
    from mutagen.mp4 import MP4

    DEPENDENCIES["audio"] = True
except ImportError:
    logger.warning("音频处理库缺失。安装mutagen以支持音频元数据: pip install mutagen")

try:
    import ffmpeg

    DEPENDENCIES["video"] = True
except ImportError:
    logger.warning("视频处理库缺失。安装ffmpeg-python以支持视频元数据: pip install ffmpeg-python")

try:
    import PyPDF2
    import docx

    DEPENDENCIES["document"] = True
except ImportError:
    logger.warning("文档处理库缺失。安装PyPDF2和python-docx以支持文档元数据: pip install PyPDF2 python-docx")


class FileType(Enum):
    """文件类型枚举"""
    UNKNOWN = "unknown"
    IMAGE = "image"
    AUDIO = "audio"
    VIDEO = "video"
    DOCUMENT = "document"


class OutputFormat(Enum):
    """输出格式枚举"""
    TEXT = "text"
    JSON = "json"
    CSV = "csv"
    XML = "xml"


class MetadataEditor:
    """元数据编辑器类，提供查看和修改文件元数据的功能"""

    # 支持的文件扩展名映射到文件类型
    FILE_EXTENSIONS = {
        # 图像文件
        ".jpg": FileType.IMAGE,
        ".jpeg": FileType.IMAGE,
        ".png": FileType.IMAGE,
        ".tiff": FileType.IMAGE,
        ".tif": FileType.IMAGE,
        ".webp": FileType.IMAGE,
        ".heif": FileType.IMAGE,
        ".heic": FileType.IMAGE,

        # 音频文件
        ".mp3": FileType.AUDIO,
        ".flac": FileType.AUDIO,
        ".wav": FileType.AUDIO,
        ".ogg": FileType.AUDIO,
        ".m4a": FileType.AUDIO,
        ".aac": FileType.AUDIO,

        # 视频文件
        ".mp4": FileType.VIDEO,
        ".mov": FileType.VIDEO,
        ".avi": FileType.VIDEO,
        ".mkv": FileType.VIDEO,
        ".webm": FileType.VIDEO,

        # 文档文件
        ".pdf": FileType.DOCUMENT,
        ".docx": FileType.DOCUMENT,
        ".xlsx": FileType.DOCUMENT,
        ".pptx": FileType.DOCUMENT,
    }

    def __init__(
            self,
            files: List[str],
            output_format: OutputFormat = OutputFormat.TEXT,
            recursive: bool = False,
            include_patterns: Optional[List[str]] = None,
            exclude_patterns: Optional[List[str]] = None,
            metadata_fields: Optional[List[str]] = None,
            add_metadata: Optional[Dict[str, str]] = None,
            remove_metadata: Optional[List[str]] = None,
            backup: bool = False,
            output_file: Optional[str] = None,
            export_file: Optional[str] = None,
            import_file: Optional[str] = None,
            preserve_original: bool = True,
            verbose: bool = False,
            dry_run: bool = False
    ):
        """
        初始化元数据编辑器
        
        Args:
            files: 要处理的文件或目录列表
            output_format: 输出格式
            recursive: 是否递归处理子目录
            include_patterns: 要包含的文件模式列表
            exclude_patterns: 要排除的文件模式列表
            metadata_fields: 要处理的元数据字段列表
            add_metadata: 要添加或修改的元数据字段和值
            remove_metadata: 要删除的元数据字段列表
            backup: 是否在修改前备份文件
            output_file: 输出文件路径
            export_file: 元数据导出文件路径
            import_file: 元数据导入文件路径
            preserve_original: 是否保留原始元数据（当添加新字段时）
            verbose: 是否显示详细信息
            dry_run: 是否仅模拟运行而不实际修改文件
        """
        self.files = files
        self.output_format = output_format
        self.recursive = recursive
        self.include_patterns = include_patterns or ["*.*"]
        self.exclude_patterns = exclude_patterns or []
        self.metadata_fields = metadata_fields
        self.add_metadata = add_metadata or {}
        self.remove_metadata = remove_metadata or []
        self.backup = backup
        self.output_file = output_file
        self.export_file = export_file
        self.import_file = import_file
        self.preserve_original = preserve_original
        self.verbose = verbose
        self.dry_run = dry_run

        # 处理统计
        self.processed_files = 0
        self.modified_files = 0
        self.error_files = 0
        self.errors = []

        # 导入元数据（如果指定了导入文件）
        self.import_data = {}
        if self.import_file:
            self._import_metadata()

    def process_files(self) -> bool:
        """
        处理文件列表
        
        Returns:
            处理是否成功
        """
        # 重置统计信息
        self.processed_files = 0
        self.modified_files = 0
        self.error_files = 0
        self.errors = []

        # 收集所有符合条件的文件
        all_files = self._collect_files()

        if not all_files:
            logger.warning("未找到匹配的文件")
            return False

        logger.info(f"找到 {len(all_files)} 个文件需要处理")

        # 用于存储所有元数据的字典
        all_metadata = {}

        # 处理每个文件
        for file_path in all_files:
            try:
                file_type = self._detect_file_type(file_path)

                if file_type == FileType.UNKNOWN:
                    logger.warning(f"无法识别文件类型: {file_path}")
                    continue

                if self.verbose:
                    logger.info(f"处理文件: {file_path} (类型: {file_type.value})")

                # 读取元数据
                metadata = self._read_metadata(file_path, file_type)

                if metadata:
                    # 存储到总元数据字典中
                    all_metadata[file_path] = metadata

                    # 如果需要修改元数据
                    if self.add_metadata or self.remove_metadata or self.import_data:
                        if not self.dry_run:
                            modified = self._modify_metadata(file_path, file_type, metadata)
                            if modified:
                                self.modified_files += 1
                        else:
                            logger.info(f"[模拟] 将修改文件: {file_path}")

                self.processed_files += 1

            except Exception as e:
                self.error_files += 1
                error_msg = f"处理文件 {file_path} 时出错: {str(e)}"
                self.errors.append(error_msg)
                logger.error(error_msg)

        # 输出或导出元数据
        if all_metadata:
            if self.export_file:
                self._export_metadata(all_metadata)

            if self.output_file:
                self._output_metadata(all_metadata)
            else:
                # 直接输出到控制台
                self._print_metadata(all_metadata)

        # 打印总结
        logger.info(f"处理完成. 总计: {self.processed_files}, 修改: {self.modified_files}, 错误: {self.error_files}")

        if self.errors:
            logger.warning("处理过程中出现以下错误:")
            for error in self.errors:
                logger.warning(f"  - {error}")

        return self.error_files == 0

    def _collect_files(self) -> List[str]:
        """
        收集所有符合条件的文件
        
        Returns:
            文件路径列表
        """
        collected_files = []

        for path in self.files:
            if os.path.isfile(path):
                # 如果是文件，直接检查是否符合条件
                if self._should_process_file(path):
                    collected_files.append(path)
            elif os.path.isdir(path):
                # 如果是目录，遍历收集文件
                for root, dirs, files in os.walk(path):
                    for filename in files:
                        file_path = os.path.join(root, filename)
                        if self._should_process_file(file_path):
                            collected_files.append(file_path)

                    # 如果不需要递归处理，则跳出循环
                    if not self.recursive:
                        break
            else:
                logger.warning(f"路径不存在或无法访问: {path}")

        return collected_files

    def _should_process_file(self, file_path: str) -> bool:
        """
        检查是否应处理该文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            是否应处理该文件
        """
        # 检查文件是否存在
        if not os.path.isfile(file_path):
            return False

        # 获取文件名
        filename = os.path.basename(file_path)

        # 检查是否符合包含模式
        included = any(self._match_pattern(filename, pattern) for pattern in self.include_patterns)
        if not included:
            return False

        # 检查是否符合排除模式
        excluded = any(self._match_pattern(filename, pattern) for pattern in self.exclude_patterns)
        if excluded:
            return False

        # 检查是否是支持的文件类型
        ext = os.path.splitext(file_path)[1].lower()
        return ext in self.FILE_EXTENSIONS

    def _match_pattern(self, filename: str, pattern: str) -> bool:
        """
        检查文件名是否匹配模式
        
        Args:
            filename: 文件名
            pattern: 匹配模式（支持通配符）
            
        Returns:
            是否匹配
        """
        import fnmatch
        return fnmatch.fnmatch(filename, pattern)

    def _detect_file_type(self, file_path: str) -> FileType:
        """
        检测文件类型
        
        Args:
            file_path: 文件路径
            
        Returns:
            文件类型
        """
        ext = os.path.splitext(file_path)[1].lower()
        return self.FILE_EXTENSIONS.get(ext, FileType.UNKNOWN)

    def _read_metadata(self, file_path: str, file_type: FileType) -> Dict[str, Any]:
        """
        读取文件元数据
        
        Args:
            file_path: 文件路径
            file_type: 文件类型
            
        Returns:
            元数据字典
        """
        metadata = {}

        try:
            if file_type == FileType.IMAGE and DEPENDENCIES["image"]:
                metadata = self._read_image_metadata(file_path)
            elif file_type == FileType.AUDIO and DEPENDENCIES["audio"]:
                metadata = self._read_audio_metadata(file_path)
            elif file_type == FileType.VIDEO and DEPENDENCIES["video"]:
                metadata = self._read_video_metadata(file_path)
            elif file_type == FileType.DOCUMENT and DEPENDENCIES["document"]:
                metadata = self._read_document_metadata(file_path)
            else:
                logger.warning(f"不支持的文件类型或缺少必要的库: {file_path}")

            # 过滤元数据字段（如果指定了特定字段）
            if self.metadata_fields and metadata:
                filtered_metadata = {}
                for field in self.metadata_fields:
                    # 支持通配符匹配
                    pattern = re.compile(fnmatch.translate(field))
                    for key, value in metadata.items():
                        if pattern.match(key):
                            filtered_metadata[key] = value
                metadata = filtered_metadata

        except Exception as e:
            logger.error(f"读取元数据时出错 ({file_path}): {str(e)}")
            raise

        return metadata

    def _read_image_metadata(self, file_path: str) -> Dict[str, Any]:
        """读取图像文件元数据"""
        metadata = {}

        try:
            # 使用PIL读取基本EXIF数据
            with Image.open(file_path) as img:
                if hasattr(img, '_getexif') and img._getexif():
                    exif_data = img._getexif()
                    if exif_data:
                        for tag_id, value in exif_data.items():
                            tag_name = ExifTags.TAGS.get(tag_id, str(tag_id))
                            # 处理特殊类型
                            if isinstance(value, bytes):
                                try:
                                    value = value.decode('utf-8')
                                except UnicodeDecodeError:
                                    value = str(value)
                            metadata[tag_name] = value

                # 读取基本图像信息
                metadata["ImageWidth"] = img.width
                metadata["ImageHeight"] = img.height
                metadata["ImageFormat"] = img.format
                metadata["ImageMode"] = img.mode

            # 使用piexif获取更多EXIF数据
            try:
                exif_dict = piexif.load(file_path)
                for ifd_name in exif_dict:
                    if ifd_name != 'thumbnail':
                        for tag, value in exif_dict[ifd_name].items():
                            tag_name = piexif.TAGS[ifd_name].get(tag, {}).get('name', str(tag))
                            # 处理特殊类型
                            if isinstance(value, bytes):
                                try:
                                    value = value.decode('utf-8')
                                except UnicodeDecodeError:
                                    value = str(value)
                            elif isinstance(value, tuple) and len(value) == 2:
                                # 分数值转换为浮点数
                                if value[1] != 0:
                                    value = value[0] / value[1]
                            metadata[tag_name] = value
            except:
                # piexif可能对某些文件格式不起作用
                pass

        except Exception as e:
            logger.error(f"读取图像元数据时出错: {str(e)}")
            raise

        return metadata

    def _read_audio_metadata(self, file_path: str) -> Dict[str, Any]:
        """读取音频文件元数据"""
        metadata = {}

        try:
            # 使用mutagen读取音频元数据
            audio = mutagen.File(file_path)
            if audio:
                # 处理基本元数据
                metadata["FileFormat"] = audio.mime[0].split('/')[-1] if hasattr(audio,
                                                                                 'mime') and audio.mime else "unknown"
                metadata["Length"] = audio.info.length if hasattr(audio.info, 'length') else 0
                metadata["Bitrate"] = audio.info.bitrate if hasattr(audio.info, 'bitrate') else 0

                # 处理特定格式的标签
                if isinstance(audio, MP3):
                    # 处理ID3标签
                    if audio.tags:
                        for key, value in audio.tags.items():
                            # 忽略专辑封面等二进制数据
                            if key.startswith('APIC:'):
                                metadata[f"{key} (Cover Image)"] = "Binary data"
                            else:
                                metadata[key] = str(value)

                elif isinstance(audio, FLAC):
                    # 处理FLAC标签
                    if audio.tags:
                        for key, value in audio.tags.items():
                            metadata[key] = ', '.join(value)

                elif isinstance(audio, MP4):
                    # 处理MP4标签
                    if audio.tags:
                        for key, value in audio.tags.items():
                            if key == 'covr':  # 专辑封面
                                metadata[f"{key} (Cover Image)"] = "Binary data"
                            else:
                                metadata[key] = str(value)
                else:
                    # 通用标签处理
                    for key, value in audio.items():
                        if isinstance(value, list):
                            metadata[key] = ', '.join(str(v) for v in value)
                        else:
                            metadata[key] = str(value)

        except Exception as e:
            logger.error(f"读取音频元数据时出错: {str(e)}")
            raise

        return metadata

    def _read_video_metadata(self, file_path: str) -> Dict[str, Any]:
        """读取视频文件元数据"""
        metadata = {}

        try:
            # 使用ffmpeg-python读取视频元数据
            probe = ffmpeg.probe(file_path)

            # 基本文件信息
            metadata["FileFormat"] = probe.get('format', {}).get('format_name', '')
            metadata["Duration"] = float(probe.get('format', {}).get('duration', 0))
            metadata["Size"] = int(probe.get('format', {}).get('size', 0))
            metadata["Bitrate"] = int(probe.get('format', {}).get('bit_rate', 0))

            # 流信息
            for stream in probe.get('streams', []):
                stream_type = stream.get('codec_type', '')
                if stream_type == 'video':
                    metadata["VideoCodec"] = stream.get('codec_name', '')
                    metadata["VideoWidth"] = stream.get('width', 0)
                    metadata["VideoHeight"] = stream.get('height', 0)
                    metadata["FrameRate"] = self._parse_frame_rate(stream.get('r_frame_rate', ''))
                elif stream_type == 'audio':
                    metadata["AudioCodec"] = stream.get('codec_name', '')
                    metadata["AudioChannels"] = stream.get('channels', 0)
                    metadata["AudioSampleRate"] = stream.get('sample_rate', '')

            # 标签信息
            tags = probe.get('format', {}).get('tags', {})
            for key, value in tags.items():
                metadata[f"Tag_{key}"] = value

        except Exception as e:
            logger.error(f"读取视频元数据时出错: {str(e)}")
            raise

        return metadata

    def _parse_frame_rate(self, rate_str: str) -> float:
        """解析帧率字符串（如 '24000/1001'）"""
        if not rate_str or '/' not in rate_str:
            return 0.0
        try:
            num, den = map(int, rate_str.split('/'))
            return num / den if den != 0 else 0.0
        except:
            return 0.0

    def _read_document_metadata(self, file_path: str) -> Dict[str, Any]:
        """读取文档文件元数据"""
        metadata = {}
        ext = os.path.splitext(file_path)[1].lower()

        try:
            if ext == '.pdf':
                # 读取PDF元数据
                with open(file_path, 'rb') as f:
                    pdf = PyPDF2.PdfReader(f)
                    # 文档信息
                    metadata["PageCount"] = len(pdf.pages)

                    # 文档属性
                    if pdf.metadata:
                        for key, value in pdf.metadata.items():
                            if key.startswith('/'):
                                key = key[1:]  # 移除前缀斜杠
                            metadata[key] = value

            elif ext == '.docx':
                # 读取Word文档元数据
                doc = docx.Document(file_path)
                core_props = doc.core_properties

                # 基本属性
                metadata["Title"] = core_props.title or ""
                metadata["Author"] = core_props.author or ""
                metadata["Subject"] = core_props.subject or ""
                metadata["Keywords"] = core_props.keywords or ""
                metadata["Comments"] = core_props.comments or ""
                metadata["LastModifiedBy"] = core_props.last_modified_by or ""
                metadata["Created"] = str(core_props.created) if core_props.created else ""
                metadata["Modified"] = str(core_props.modified) if core_props.modified else ""
                metadata["Category"] = core_props.category or ""

                # 文档统计
                metadata["Paragraphs"] = len(doc.paragraphs)
                metadata["Tables"] = len(doc.tables)

        except Exception as e:
            logger.error(f"读取文档元数据时出错: {str(e)}")
            raise

        return metadata

    def _modify_metadata(self, file_path: str, file_type: FileType, current_metadata: Dict[str, Any]) -> bool:
        """
        修改文件元数据
        
        Args:
            file_path: 文件路径
            file_type: 文件类型
            current_metadata: 当前元数据
            
        Returns:
            是否成功修改
        """
        # 备份文件（如果启用）
        if self.backup:
            backup_path = f"{file_path}.bak"
            try:
                import shutil
                shutil.copy2(file_path, backup_path)
                if self.verbose:
                    logger.info(f"已创建备份: {backup_path}")
            except Exception as e:
                logger.error(f"创建备份时出错: {str(e)}")
                return False

        # 获取要应用的元数据
        metadata_to_apply = {}

        # 从导入数据中获取元数据（如果有）
        if self.import_data and file_path in self.import_data:
            metadata_to_apply.update(self.import_data[file_path])

        # 从命令行参数中获取元数据
        if self.add_metadata:
            metadata_to_apply.update(self.add_metadata)

        # 如果没有要应用的元数据和要删除的字段，则返回
        if not metadata_to_apply and not self.remove_metadata:
            return False

        try:
            if file_type == FileType.IMAGE and DEPENDENCIES["image"]:
                return self._modify_image_metadata(file_path, metadata_to_apply, self.remove_metadata)
            elif file_type == FileType.AUDIO and DEPENDENCIES["audio"]:
                return self._modify_audio_metadata(file_path, metadata_to_apply, self.remove_metadata)
            elif file_type == FileType.VIDEO and DEPENDENCIES["video"]:
                return self._modify_video_metadata(file_path, metadata_to_apply, self.remove_metadata)
            elif file_type == FileType.DOCUMENT and DEPENDENCIES["document"]:
                return self._modify_document_metadata(file_path, metadata_to_apply, self.remove_metadata)
            else:
                logger.warning(f"不支持修改此类型文件的元数据: {file_path}")
                return False

        except Exception as e:
            logger.error(f"修改元数据时出错 ({file_path}): {str(e)}")
            return False

    def _modify_image_metadata(self, file_path: str, metadata: Dict[str, Any], fields_to_remove: List[str]) -> bool:
        """修改图像文件元数据"""
        try:
            # 尝试使用piexif修改EXIF数据（对JPEG文件最有效）
            ext = os.path.splitext(file_path)[1].lower()
            if ext in ['.jpg', '.jpeg'] and piexif:
                # 读取现有EXIF数据
                try:
                    exif_dict = piexif.load(file_path)
                except:
                    # 如果没有EXIF数据，创建新的
                    exif_dict = {'0th': {}, '1st': {}, 'Exif': {}, 'GPS': {}, 'Interop': {}}

                # 修改字段
                modified = False

                # 删除字段
                for field in fields_to_remove:
                    for ifd in exif_dict:
                        if ifd == 'thumbnail':
                            continue

                        # 尝试在piexif中找到标签ID
                        tag_id = None
                        for tag, tag_info in piexif.TAGS[ifd].items():
                            if tag_info['name'] == field:
                                tag_id = tag
                                break

                        if tag_id and tag_id in exif_dict[ifd]:
                            del exif_dict[ifd][tag_id]
                            modified = True
                            if self.verbose:
                                logger.info(f"已删除EXIF字段: {field}")

                # 添加/修改字段
                for field, value in metadata.items():
                    # 找到适当的IFD（Image File Directory）
                    target_ifd = None
                    tag_id = None

                    # 搜索所有可能的标签
                    for ifd in ['0th', '1st', 'Exif', 'GPS', 'Interop']:
                        for tag, tag_info in piexif.TAGS[ifd].items():
                            if tag_info['name'] == field:
                                target_ifd = ifd
                                tag_id = tag
                                break
                        if target_ifd:
                            break

                    if not target_ifd:
                        # 如果找不到标签，默认使用Exif IFD
                        logger.warning(f"未找到EXIF标签: {field}，跳过")
                        continue

                    # 转换值为适当的格式
                    try:
                        # 对于特定类型的字段进行特殊处理
                        if isinstance(value, (int, float)):
                            # 一些数值需要存储为分数
                            if field in ['FNumber', 'ExposureTime', 'ApertureValue']:
                                # 存储为分数 (numerator, denominator)
                                if value == 0:
                                    exif_value = (0, 1)
                                else:
                                    # 简化分数表示
                                    from fractions import Fraction
                                    frac = Fraction(value).limit_denominator(1000)
                                    exif_value = (frac.numerator, frac.denominator)
                            else:
                                exif_value = int(value)
                        elif isinstance(value, str):
                            exif_value = value.encode('utf-8')
                        elif isinstance(value, bytes):
                            exif_value = value
                        elif isinstance(value, (list, tuple)) and len(value) == 2:
                            # 假设是分数值
                            exif_value = (int(value[0]), int(value[1]))
                        else:
                            # 转换为字符串
                            exif_value = str(value).encode('utf-8')

                        # 设置值
                        exif_dict[target_ifd][tag_id] = exif_value
                        modified = True
                        if self.verbose:
                            logger.info(f"已设置EXIF字段: {field} = {value}")

                    except Exception as e:
                        logger.error(f"设置EXIF字段 {field} 时出错: {str(e)}")

                # 如果有修改，保存回文件
                if modified:
                    exif_bytes = piexif.dump(exif_dict)
                    piexif.insert(exif_bytes, file_path)
                    return True

            # 使用PIL修改基本属性（适用于支持的其他图像格式）
            with Image.open(file_path) as img:
                # 目前PIL对修改元数据支持有限
                # 这里可以实现特定格式的扩展
                pass

            return False

        except Exception as e:
            logger.error(f"修改图像元数据时出错: {str(e)}")
            raise

    def _modify_audio_metadata(self, file_path: str, metadata: Dict[str, Any], fields_to_remove: List[str]) -> bool:
        """修改音频文件元数据"""
        try:
            # 使用mutagen修改音频元数据
            audio = mutagen.File(file_path)
            if not audio:
                logger.warning(f"无法读取音频文件: {file_path}")
                return False

            modified = False

            # 特定格式的处理
            if isinstance(audio, MP3):
                # 确保文件有ID3标签
                if not audio.tags:
                    audio.tags = ID3()

                # 删除字段
                for field in fields_to_remove:
                    if field in audio.tags:
                        del audio.tags[field]
                        modified = True
                        if self.verbose:
                            logger.info(f"已删除ID3标签: {field}")

                # 添加/修改字段
                for field, value in metadata.items():
                    # ID3有特定的帧类型
                    if field.lower() == "title":
                        audio.tags.add(TIT2(text=[value]))
                        modified = True
                    elif field.lower() == "artist":
                        audio.tags.add(TPE1(text=[value]))
                        modified = True
                    elif field.lower() == "album":
                        audio.tags.add(TALB(text=[value]))
                        modified = True
                    elif field.lower() == "year":
                        audio.tags.add(TDRC(text=[value]))
                        modified = True
                    else:
                        # 对于其他标签，尝试直接设置（可能不总是有效）
                        try:
                            if field in audio.tags:
                                audio.tags[field] = value
                                modified = True
                        except:
                            logger.warning(f"无法设置ID3标签: {field}")

            elif isinstance(audio, FLAC):
                # 删除字段
                for field in fields_to_remove:
                    if field in audio:
                        del audio[field]
                        modified = True
                        if self.verbose:
                            logger.info(f"已删除FLAC标签: {field}")

                # 添加/修改字段
                for field, value in metadata.items():
                    # FLAC使用简单的键值对
                    audio[field] = value
                    modified = True
                    if self.verbose:
                        logger.info(f"已设置FLAC标签: {field} = {value}")

            elif isinstance(audio, MP4):
                # 删除字段
                for field in fields_to_remove:
                    if field in audio:
                        del audio[field]
                        modified = True
                        if self.verbose:
                            logger.info(f"已删除MP4标签: {field}")

                # 添加/修改字段
                for field, value in metadata.items():
                    # MP4有特定的原子名称
                    # 简化处理，可能需要扩展
                    audio[field] = value
                    modified = True
                    if self.verbose:
                        logger.info(f"已设置MP4标签: {field} = {value}")
            else:
                # 通用处理
                # 删除字段
                for field in fields_to_remove:
                    if field in audio:
                        del audio[field]
                        modified = True
                        if self.verbose:
                            logger.info(f"已删除标签: {field}")

                # 添加/修改字段
                for field, value in metadata.items():
                    audio[field] = value
                    modified = True
                    if self.verbose:
                        logger.info(f"已设置标签: {field} = {value}")

            # 保存更改
            if modified:
                audio.save()
                return True

            return False

        except Exception as e:
            logger.error(f"修改音频元数据时出错: {str(e)}")
            raise

    def _modify_video_metadata(self, file_path: str, metadata: Dict[str, Any], fields_to_remove: List[str]) -> bool:
        """修改视频文件元数据"""
        try:
            # 注意: 修改视频元数据可能需要重新编码，这里只实现基本功能
            # 使用ffmpeg添加/修改元数据

            # 构建元数据参数
            metadata_args = {}
            for field, value in metadata.items():
                metadata_args[f"metadata:{field}"] = str(value)

            # 删除字段 (通过设置为空值)
            for field in fields_to_remove:
                metadata_args[f"metadata:{field}"] = ""

            if not metadata_args:
                return False

            # 创建临时文件
            import tempfile
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file_path)[1])
            temp_file.close()

            try:
                # 使用ffmpeg设置元数据
                input_stream = ffmpeg.input(file_path)
                output_stream = ffmpeg.output(input_stream, temp_file.name, **metadata_args, codec="copy")
                ffmpeg.run(output_stream, quiet=not self.verbose, overwrite_output=True)

                # 替换原文件
                import shutil
                shutil.move(temp_file.name, file_path)

                return True

            finally:
                # 清理临时文件（如果还存在）
                if os.path.exists(temp_file.name):
                    os.unlink(temp_file.name)

        except Exception as e:
            logger.error(f"修改视频元数据时出错: {str(e)}")
            raise

    def _modify_document_metadata(self, file_path: str, metadata: Dict[str, Any], fields_to_remove: List[str]) -> bool:
        """修改文档文件元数据"""
        ext = os.path.splitext(file_path)[1].lower()

        try:
            if ext == '.pdf':
                # 修改PDF元数据
                with open(file_path, 'rb') as f:
                    pdf = PyPDF2.PdfReader(f)
                    writer = PyPDF2.PdfWriter()

                    # 复制所有页面
                    for page in pdf.pages:
                        writer.add_page(page)

                    # 获取现有元数据
                    current_metadata = pdf.metadata or {}

                    # 创建新的元数据字典
                    new_metadata = {}

                    # 复制现有元数据（排除要删除的字段）
                    for key, value in current_metadata.items():
                        # 去除前缀斜杠
                        clean_key = key[1:] if key.startswith('/') else key
                        if clean_key not in fields_to_remove:
                            new_metadata[f"/{clean_key}"] = value

                    # 添加新的元数据
                    for key, value in metadata.items():
                        # 确保键有前缀斜杠
                        if not key.startswith('/'):
                            key = f"/{key}"
                        new_metadata[key] = value

                    # 设置元数据
                    writer.add_metadata(new_metadata)

                    # 写入临时文件
                    import tempfile
                    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
                    temp_file.close()

                    with open(temp_file.name, 'wb') as f:
                        writer.write(f)

                    # 替换原文件
                    import shutil
                    shutil.move(temp_file.name, file_path)

                    return True

            elif ext == '.docx':
                # 修改Word文档元数据
                doc = docx.Document(file_path)
                core_props = doc.core_properties
                modified = False

                # 处理标准属性
                property_map = {
                    "Title": "title",
                    "Author": "author",
                    "Subject": "subject",
                    "Keywords": "keywords",
                    "Comments": "comments",
                    "Category": "category"
                }

                # 添加/修改属性
                for key, value in metadata.items():
                    prop_name = property_map.get(key)
                    if prop_name and hasattr(core_props, prop_name):
                        setattr(core_props, prop_name, value)
                        modified = True
                        if self.verbose:
                            logger.info(f"已设置文档属性: {key} = {value}")

                # 保存文档
                if modified:
                    doc.save(file_path)
                    return True

            return False

        except Exception as e:
            logger.error(f"修改文档元数据时出错: {str(e)}")
            raise

    def _import_metadata(self):
        """从文件导入元数据"""
        if not os.path.isfile(self.import_file):
            logger.error(f"导入文件不存在: {self.import_file}")
            return

        ext = os.path.splitext(self.import_file)[1].lower()

        try:
            if ext == '.json':
                # 从JSON文件导入
                with open(self.import_file, 'r', encoding='utf-8') as f:
                    self.import_data = json.load(f)

            elif ext == '.csv':
                # 从CSV文件导入
                self.import_data = {}
                with open(self.import_file, 'r', encoding='utf-8') as f:
                    reader = csv.reader(f)
                    headers = next(reader)  # 第一行是表头

                    # 确保第一列是文件路径
                    if len(headers) < 2:
                        logger.error(f"CSV文件格式无效: {self.import_file}")
                        return

                    file_col = 0  # 假设第一列是文件路径

                    for row in reader:
                        if len(row) < len(headers):
                            continue  # 跳过不完整的行

                        file_path = row[file_col]
                        file_metadata = {}

                        # 解析每个字段
                        for i in range(1, len(headers)):
                            if i < len(row) and row[i]:  # 确保值不为空
                                file_metadata[headers[i]] = row[i]

                        if file_metadata:
                            self.import_data[file_path] = file_metadata
            else:
                logger.error(f"不支持的导入文件格式: {ext}")

            logger.info(f"已从 {self.import_file} 导入元数据，共 {len(self.import_data)} 个文件")

        except Exception as e:
            logger.error(f"导入元数据时出错: {str(e)}")
            self.import_data = {}

    def _export_metadata(self, metadata_dict: Dict[str, Dict[str, Any]]):
        """导出元数据到文件"""
        if not self.export_file:
            return

        ext = os.path.splitext(self.export_file)[1].lower()

        try:
            if ext == '.json':
                # 转换不可序列化的对象为字符串
                serializable_metadata = {}
                for file_path, metadata in metadata_dict.items():
                    serializable_metadata[file_path] = {}
                    for key, value in metadata.items():
                        if isinstance(value, (str, int, float, bool, type(None))):
                            serializable_metadata[file_path][key] = value
                        else:
                            serializable_metadata[file_path][key] = str(value)

                # 导出为JSON
                with open(self.export_file, 'w', encoding='utf-8') as f:
                    json.dump(serializable_metadata, f, ensure_ascii=False, indent=2)

            elif ext == '.csv':
                # 收集所有可能的字段
                all_fields = set()
                for metadata in metadata_dict.values():
                    all_fields.update(metadata.keys())

                # 排序字段
                sorted_fields = sorted(all_fields)

                # 导出为CSV
                with open(self.export_file, 'w', encoding='utf-8', newline='') as f:
                    writer = csv.writer(f)

                    # 写入表头
                    writer.writerow(['FilePath'] + sorted_fields)

                    # 写入每个文件的数据
                    for file_path, metadata in metadata_dict.items():
                        row = [file_path]
                        for field in sorted_fields:
                            value = metadata.get(field, '')
                            # 确保值是字符串
                            if not isinstance(value, str):
                                value = str(value)
                            row.append(value)
                        writer.writerow(row)

            elif ext == '.xml':
                # 导出为XML
                import xml.dom.minidom as md

                doc = md.getDOMImplementation().createDocument(None, "Metadata", None)
                root = doc.documentElement

                for file_path, metadata in metadata_dict.items():
                    file_elem = doc.createElement("File")
                    file_elem.setAttribute("path", file_path)

                    for key, value in metadata.items():
                        if isinstance(value, (str, int, float, bool, type(None))):
                            field_elem = doc.createElement("Field")
                            field_elem.setAttribute("name", key)

                            # 转换不同类型的值
                            if value is None:
                                text_value = ""
                            else:
                                text_value = str(value)

                            text_node = doc.createTextNode(text_value)
                            field_elem.appendChild(text_node)
                            file_elem.appendChild(field_elem)

                    root.appendChild(file_elem)

                # 写入文件
                with open(self.export_file, 'w', encoding='utf-8') as f:
                    f.write(doc.toprettyxml(indent="  "))
            else:
                logger.error(f"不支持的导出文件格式: {ext}")
                return

            logger.info(f"已导出元数据到 {self.export_file}")

        except Exception as e:
            logger.error(f"导出元数据时出错: {str(e)}")

    def _output_metadata(self, metadata_dict: Dict[str, Dict[str, Any]]):
        """输出元数据到文件"""
        if not self.output_file:
            return

        try:
            with open(self.output_file, 'w', encoding='utf-8') as f:
                if self.output_format == OutputFormat.TEXT:
                    for file_path, metadata in metadata_dict.items():
                        f.write(f"文件: {file_path}\n")
                        f.write("-" * 80 + "\n")

                        for key, value in metadata.items():
                            f.write(f"{key}: {value}\n")

                        f.write("\n\n")

                elif self.output_format == OutputFormat.JSON:
                    # 转换不可序列化的对象为字符串
                    serializable_metadata = {}
                    for file_path, metadata in metadata_dict.items():
                        serializable_metadata[file_path] = {}
                        for key, value in metadata.items():
                            if isinstance(value, (str, int, float, bool, type(None))):
                                serializable_metadata[file_path][key] = value
                            else:
                                serializable_metadata[file_path][key] = str(value)

                    json.dump(serializable_metadata, f, ensure_ascii=False, indent=2)

                elif self.output_format == OutputFormat.CSV:
                    # 收集所有可能的字段
                    all_fields = set()
                    for metadata in metadata_dict.values():
                        all_fields.update(metadata.keys())

                    # 排序字段
                    sorted_fields = sorted(all_fields)

                    # 创建CSV写入器
                    writer = csv.writer(f)

                    # 写入表头
                    writer.writerow(['FilePath'] + sorted_fields)

                    # 写入每个文件的数据
                    for file_path, metadata in metadata_dict.items():
                        row = [file_path]
                        for field in sorted_fields:
                            value = metadata.get(field, '')
                            # 确保值是字符串
                            if not isinstance(value, str):
                                value = str(value)
                            row.append(value)
                        writer.writerow(row)

                elif self.output_format == OutputFormat.XML:
                    # 创建XML
                    import xml.dom.minidom as md

                    doc = md.getDOMImplementation().createDocument(None, "Metadata", None)
                    root = doc.documentElement

                    for file_path, metadata in metadata_dict.items():
                        file_elem = doc.createElement("File")
                        file_elem.setAttribute("path", file_path)

                        for key, value in metadata.items():
                            if isinstance(value, (str, int, float, bool, type(None))):
                                field_elem = doc.createElement("Field")
                                field_elem.setAttribute("name", key)

                                # 转换不同类型的值
                                if value is None:
                                    text_value = ""
                                else:
                                    text_value = str(value)

                                text_node = doc.createTextNode(text_value)
                                field_elem.appendChild(text_node)
                                file_elem.appendChild(field_elem)

                        root.appendChild(file_elem)

                    # 写入文件
                    f.write(doc.toprettyxml(indent="  "))

            logger.info(f"已输出元数据到 {self.output_file}")

        except Exception as e:
            logger.error(f"输出元数据时出错: {str(e)}")

    def _print_metadata(self, metadata_dict: Dict[str, Dict[str, Any]]):
        """打印元数据到控制台"""
        for file_path, metadata in metadata_dict.items():
            print(f"\n文件: {file_path}")
            print("-" * 80)

            # 格式化打印
            max_key_length = max([len(key) for key in metadata.keys()]) if metadata else 0

            for key, value in sorted(metadata.items()):
                # 格式化输出，对齐键
                print(f"{key.ljust(max_key_length)}: {value}")


def parse_args():
    """解析命令行参数"""
    parser = argparse.ArgumentParser(description='文件元数据编辑器 - 查看和修改各种文件类型的元数据')

    # 文件选择参数
    parser.add_argument('files', nargs='+', help='要处理的文件或目录')
    parser.add_argument('-r', '--recursive', action='store_true', help='递归处理子目录')
    parser.add_argument('--include', nargs='+', help='包含的文件模式（例如"*.jpg *.png"）')
    parser.add_argument('--exclude', nargs='+', help='排除的文件模式（例如"*thumb*"）')

    # 元数据选择参数
    parser.add_argument('--fields', nargs='+', help='要显示的元数据字段（支持通配符）')

    # 修改参数
    parser.add_argument('--add', nargs='+', help='要添加/修改的元数据字段和值，格式为"字段=值"')
    parser.add_argument('--remove', nargs='+', help='要删除的元数据字段')
    parser.add_argument('--preserve', action='store_true', help='保留原始元数据（当添加新字段时）')

    # 导入/导出参数
    parser.add_argument('--import-file', help='从文件导入元数据 (JSON/CSV)')
    parser.add_argument('--export-file', help='导出元数据到文件 (JSON/CSV/XML)')

    # 输出参数
    parser.add_argument('-o', '--output', help='输出结果到文件')
    parser.add_argument('--format', choices=['text', 'json', 'csv', 'xml'], default='text',
                        help='输出格式 (默认: text)')

    # 其他参数
    parser.add_argument('--backup', action='store_true', help='在修改前备份文件')
    parser.add_argument('-v', '--verbose', action='store_true', help='显示详细信息')
    parser.add_argument('--dry-run', action='store_true', help='模拟运行，不实际修改文件')

    return parser.parse_args()


def main():
    """主函数"""
    args = parse_args()

    # 准备元数据字段
    add_metadata = {}
    if args.add:
        for field_value in args.add:
            if '=' in field_value:
                field, value = field_value.split('=', 1)
                add_metadata[field] = value
            else:
                logger.warning(f"忽略无效的元数据字段格式: {field_value}，应为'字段=值'")

    # 创建编辑器实例
    editor = MetadataEditor(
        files=args.files,
        output_format=OutputFormat[args.format.upper()],
        recursive=args.recursive,
        include_patterns=args.include,
        exclude_patterns=args.exclude,
        metadata_fields=args.fields,
        add_metadata=add_metadata,
        remove_metadata=args.remove,
        backup=args.backup,
        output_file=args.output,
        export_file=args.export_file,
        import_file=args.import_file,
        preserve_original=args.preserve,
        verbose=args.verbose,
        dry_run=args.dry_run
    )

    # 处理文件
    success = editor.process_files()

    # 返回退出码
    return 0 if success else 1


if __name__ == '__main__':
    sys.exit(main())
